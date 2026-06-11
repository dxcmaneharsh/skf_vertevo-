import sys
import subprocess
import os

# Install python-docx if not available
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('SKF Forecast Model - Data Dictionary', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('This document details the meaning of every column in the skf_sample_data_v3_variations dataset and its specific use case for the SAP RPT/GenAI Forecasting Model.')

# Section 1
doc.add_heading('Primary Identifiers & Target Metric', level=1)
doc.add_paragraph('These columns form the core structure of what the AI is trying to predict.')

data = [
    ('Month', 'The YYYY-MM time block the row represents (Jan 2025 – Dec 2025).', 'Serves as the primary time series index. The model uses this to establish baseline seasonality, trendlines, and month-over-month (MoM) growth rates.'),
    ('OEM (Original Equipment Manufacturer)', 'The specific automotive manufacturer purchasing the bearings (e.g., OEM_A_Global, OEM_B_Euro, OEM_C_Eco).', 'A categorical variation dimension. It allows the model to learn that different regional customers buy in vastly different volumes and have different baseline needs.'),
    ('Powertrain', 'The vehicle propulsion system the bearings are being ordered for (ICE, EV, Hybrid).', 'A critical variation dimension. It allows the AI to learn product mix shifts. For example, the model learns that EV pipelines favor Miniature/Ceramic bearings over Thrust bearings.'),
    ('Product_Category', 'The specific type of SKF ball bearing being sold.', 'Acts as the dimension dividing the portfolio. The model uses this to generate independent, product-specific forecasts.'),
    ('Sales_Volume', 'The actual number of units sold for that specific row\'s combination.', 'This is the Target Variable (Y)! This is the exact number the GenAI model is trying to predict accurately.')
]

for name, meaning, use_case in data:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    p.add_run(f'\nMeaning: ').italic = True
    p.add_run(meaning)
    p.add_run(f'\nUse Case in the Forecast Model: ').italic = True
    p.add_run(use_case)

doc.add_heading('Independent Causal Regressors (External Drivers)', level=1)
doc.add_paragraph('These are the most important columns for an "Iterative Agentic / Causal Forecast" demo. These teach the model why demand shifts, rather than just guessing based on past trends.')

data2 = [
    ('Light_Vehicle_Production_Index', 'A baseline metric indicating how many total cars are being built by OEMs globally.', 'Used primarily to predict high-volume, generic bearings. It tells the model: "If auto manufacturing goes up globally this month, Deep Groove bearing sales must go up too."'),
    ('EV_Hybrid_Penetration_Pct', 'The percentage of the automotive market shifting away from traditional ICE vehicles toward electric.', 'A technology-shift regressor. The model uses this to forecast the decline of legacy components (Thrust Bearings) and aggressive spikes in new-tech components (Angular Contact & Ceramic Bearings).'),
    ('GDP_Growth_Pct & Interest_Rate_Pct', 'Core macroeconomic health indicators indicating consumer spending power and auto-loan financing affordability.', 'Used as leading indicators for the entire passenger-vehicle market. If interest rates rise, the model learns to predict a downstream overall market dip a few months later, suppressing total orders.'),
    ('Steel_Price_Index', 'A raw material cost tracker for automotive steel.', 'Margin & Delay driver. In our data, the model will learn that when this index surges suddenly above normal limits, OEMs delay bulk orders to wait for prices to fall.'),
    ('Energy_Logistics_Index', 'Tracks global freight, shipping, and factory energy costs.', 'Another overhead cost driver that impacts whether SKF customers manufacture locally or consolidate long-term shipments.'),
    ('Emission_Norms_Status', 'A qualitative text column that indicates if a new major regulatory norm is active (Pre-Euro 7 vs Euro 7 Transition).', 'Regulatory event trigger. The model will read the string "Euro 7 Transition" and instantly correlate it to a sudden spike in Self-Aligning Bearings (as emissions standards often force OEMs to redesign suspension platforms).'),
    ('OEM_Production_Status', 'A qualitative text column flagging major assembly line events (Normal Production vs Platform Upgrade).', 'Event trigger. Normal time-series models can\'t predict random spikes. By feeding this column, the AI learns that a "Platform Upgrade" event causes a massive, one-time pipeline fill order.')
]

for name, meaning, use_case in data2:
    p = doc.add_paragraph()
    p.add_run(name).bold = True
    p.add_run(f'\nMeaning: ').italic = True
    p.add_run(meaning)
    p.add_run(f'\nUse Case in the Forecast Model: ').italic = True
    p.add_run(use_case)

output_file = 'SKF_Data_Dictionary.docx'
doc.save(output_file)
print(f"Successfully created {output_file}")
